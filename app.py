import streamlit as st
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Get API endpoint from environment variable
LLAMA_API_ENDPOINT = os.getenv('LLAMA_API_ENDPOINT')
if not LLAMA_API_ENDPOINT:
    st.error("Please set the LLAMA_API_ENDPOINT in your .env file")

def generate_llama_response(prompt):
    """Generate response using Llama API."""
    try:
        if not LLAMA_API_ENDPOINT:
            return "Error: API endpoint not configured"

        headers = {
            "Content-Type": "application/json",
        }
        
        data = {
            "prompt": f"""You are a professional resume writer. 
            Please enhance the following content to be more impactful and professional. 
            Use action verbs and quantifiable achievements where possible.
            
            Content to enhance: {prompt}""",
            "max_tokens": 500,
            "temperature": 0.7
        }
        
        response = requests.post(
            LLAMA_API_ENDPOINT,
            headers=headers,
            json=data
        )
        
        if response.status_code == 200:
            return response.json()["generated_text"]
        else:
            return f"Error: {response.status_code}"
            
    except Exception as e:
        return f"Error: {str(e)}"

def enhance_content(content, section):
    """Use Llama to enhance the content for a specific resume section."""
    prompt = f"Enhance this {section} content for a resume: {content}"
    return generate_llama_response(prompt)

def create_resume(data):
    """Create a professional resume document."""
    doc = Document()
    
    # Header
    name = doc.add_paragraph()
    name_run = name.add_run(data['name'])
    name_run.bold = True
    name_run.font.size = Pt(16)
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(f"{data['email']} | {data['phone']} | {data['location']}")
    
    # Sections
    sections = [
        ('Professional Summary', data['summary']),
        ('Work Experience', data['experience']),
        ('Education', data['education']),
        ('Skills', data['skills'])
    ]
    
    for title, content in sections:
        doc.add_heading(title, level=1)
        doc.add_paragraph(content)
    
    # Save document
    doc.save('resume.docx')
    return 'resume.docx'

def main():
    st.title("AI-Powered Resume Builder (Llama 3)")
    st.write("Build your professional resume with AI assistance")
    
    # Personal Information
    st.header("Personal Information")
    name = st.text_input("Full Name")
    email = st.text_input("Email")
    phone = st.text_input("Phone")
    location = st.text_input("Location")
    
    # Professional Summary
    st.header("Professional Summary")
    summary = st.text_area("Write a brief professional summary")
    if summary and st.button("Enhance Summary"):
        with st.spinner("Enhancing summary..."):
            enhanced_summary = enhance_content(summary, "professional summary")
            st.write("Enhanced Summary:")
            st.write(enhanced_summary)
            summary = enhanced_summary
    
    # Work Experience
    st.header("Work Experience")
    experience = st.text_area("Enter your work experience (separate entries with new lines)")
    if experience and st.button("Enhance Experience"):
        with st.spinner("Enhancing experience..."):
            enhanced_exp = enhance_content(experience, "work experience")
            st.write("Enhanced Experience:")
            st.write(enhanced_exp)
            experience = enhanced_exp
    
    # Education
    st.header("Education")
    education = st.text_area("Enter your education details")
    
    # Skills
    st.header("Skills")
    skills = st.text_area("Enter your skills (separate with commas)")
    
    if st.button("Generate Resume"):
        if name and email and summary and experience:
            resume_data = {
                'name': name,
                'email': email,
                'phone': phone,
                'location': location,
                'summary': summary,
                'experience': experience,
                'education': education,
                'skills': skills
            }
            
            file_path = create_resume(resume_data)
            
            with open(file_path, 'rb') as file:
                st.download_button(
                    label="Download Resume",
                    data=file,
                    file_name="resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.error("Please fill in all required fields (Name, Email, Summary, and Experience)")

if __name__ == "__main__":
    main() 